"""
CORRECTED Translation Agent - Handles DOCX and XLSX files properly
"""

import asyncio
import logging
from pathlib import Path
from typing import Dict, Optional
import io
from docx import Document
import pandas as pd
import openpyxl
from PyPDF2 import PdfReader

logger = logging.getLogger(__name__)

class TranslationAgent:
    def __init__(self, gemini_client):
        self.gemini_client = gemini_client
        logger.info("Translation Agent initialized")

    async def process_document(
        self, 
        document_path: str, 
        source_language: str = "ja", 
        target_language: str = "en"
    ) -> Dict:
        """Process document with proper file type handling"""
        try:
            document_path = Path(document_path)
            
            # Extract text based on file extension
            if document_path.suffix.lower() == '.docx':
                extracted_text = self._extract_docx_text(document_path)
            elif document_path.suffix.lower() in ['.xlsx', '.xls']:
                extracted_text = self._extract_excel_text(document_path)
            elif document_path.suffix.lower() == '.pdf':
                extracted_text = self._extract_pdf_text(document_path)
            elif document_path.suffix.lower() in ['.txt']:
                extracted_text = self._extract_text_file(document_path)
            else:
                raise ValueError(f"Unsupported file format: {document_path.suffix}")
            
            if not extracted_text.strip():
                raise Exception("No text could be extracted from the document")
            
            # Translate if needed
            if source_language != target_language:
                translated_text = await self._translate_text(extracted_text, source_language, target_language)
            else:
                translated_text = extracted_text
            
            return {
                'original_text': extracted_text,
                'translated_content': translated_text,
                'source_language': source_language,
                'target_language': target_language,
                'document_path': str(document_path)
            }
            
        except Exception as e:
            logger.error(f"Translation failed: {str(e)}")
            raise

    def _extract_docx_text(self, file_path: Path) -> str:
        """Extract text from DOCX file"""
        try:
            # Read as binary and use python-docx
            doc = Document(file_path)
            
            full_text = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    full_text.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            full_text.append(cell.text)
            
            return '\n'.join(full_text)
            
        except Exception as e:
            logger.error(f"Error extracting DOCX text: {e}")
            raise Exception(f"Failed to extract text from DOCX: {str(e)}")

    def _extract_excel_text(self, file_path: Path) -> str:
        """Extract text from Excel file (XLSX/XLS)"""
        try:
            # Use openpyxl for XLSX files
            if file_path.suffix.lower() == '.xlsx':
                workbook = openpyxl.load_workbook(file_path)
            else:
                # For older XLS files, use pandas
                df = pd.read_excel(file_path, sheet_name=None)  # Read all sheets
                text_parts = []
                for sheet_name, sheet_df in df.items():
                    text_parts.append(f"Sheet: {sheet_name}")
                    text_parts.append(sheet_df.to_string())
                return '\n'.join(text_parts)
            
            text_parts = []
            
            for sheet_name in workbook.sheetnames:
                worksheet = workbook[sheet_name]
                text_parts.append(f"Sheet: {sheet_name}")
                
                for row in worksheet.iter_rows():
                    row_data = []
                    for cell in row:
                        if cell.value is not None:
                            row_data.append(str(cell.value))
                    if row_data:
                        text_parts.append('\t'.join(row_data))
            
            return '\n'.join(text_parts)
            
        except Exception as e:
            logger.error(f"Error extracting Excel text: {e}")
            raise Exception(f"Failed to extract text from Excel: {str(e)}")

    def _extract_pdf_text(self, file_path: Path) -> str:
        """Extract text from PDF file"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PdfReader(file)
                text_parts = []
                
                for page_num, page in enumerate(pdf_reader.pages):
                    text = page.extract_text()
                    if text.strip():
                        text_parts.append(f"Page {page_num + 1}:")
                        text_parts.append(text)
                
                return '\n'.join(text_parts)
                
        except Exception as e:
            logger.error(f"Error extracting PDF text: {e}")
            raise Exception(f"Failed to extract text from PDF: {str(e)}")

    def _extract_text_file(self, file_path: Path) -> str:
        """Extract text from plain text file with encoding detection"""
        try:
            # Try UTF-8 first
            try:
                with open(file_path, 'r', encoding='utf-8') as file:
                    return file.read()
            except UnicodeDecodeError:
                # Try other encodings
                encodings = ['cp1252', 'iso-8859-1', 'latin-1']
                for encoding in encodings:
                    try:
                        with open(file_path, 'r', encoding=encoding) as file:
                            return file.read()
                    except UnicodeDecodeError:
                        continue
                
                # If all fail, read as binary and decode with error handling
                with open(file_path, 'rb') as file:
                    return file.read().decode('utf-8', errors='replace')
                    
        except Exception as e:
            logger.error(f"Error extracting text file: {e}")
            raise Exception(f"Failed to extract text from file: {str(e)}")

    async def _translate_text(self, text: str, source_lang: str, target_lang: str) -> str:
                """Translate text using Gemini API"""
                try:
                    prompt = f"""
                    Please translate the following text from {source_lang} to {target_lang}.
            
                    Text to translate:
                    {text}
            
                    Please provide only the translation without any additional commentary.
                    """
                    translated = await self.gemini_client.generate_content_async(prompt)
                    return translated
                except Exception as e:
                    logger.error(f"Translation API failed: {e}")
                    # Return original text if translation fails
                    return text
            

    # Handle FastAPI UploadFile objects
    async def process_upload_file(
        self,
        upload_file,
        source_language: str = "ja",
        target_language: str = "en"
    ) -> Dict:
        """Process UploadFile object directly without saving to disk"""
        try:
            file_content = await upload_file.read()
            filename = upload_file.filename
            
            # Determine file type and extract text
            if filename.lower().endswith('.docx'):
                extracted_text = self._extract_docx_from_bytes(file_content)
            elif filename.lower().endswith('.xlsx'):
                extracted_text = self._extract_excel_from_bytes(file_content)
            elif filename.lower().endswith('.pdf'):
                extracted_text = self._extract_pdf_from_bytes(file_content)
            elif filename.lower().endswith('.txt'):
                extracted_text = file_content.decode('utf-8', errors='replace')
            else:
                raise ValueError(f"Unsupported file format: {filename}")
            
            if not extracted_text.strip():
                raise Exception("No text could be extracted from the document")
            
            # Translate if needed
            if source_language != target_language:
                translated_text = await self._translate_text(extracted_text, source_language, target_language)
            else:
                translated_text = extracted_text
            
            return {
                'original_text': extracted_text,
                'translated_content': translated_text,
                'source_language': source_language,
                'target_language': target_language,
                'filename': filename
            }
            
        except Exception as e:
            logger.error(f"Upload file processing failed: {str(e)}")
            raise

    def _extract_docx_from_bytes(self, file_bytes: bytes) -> str:
        """Extract text from DOCX bytes"""
        try:
            doc = Document(io.BytesIO(file_bytes))
            
            full_text = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    full_text.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            full_text.append(cell.text)
            
            return '\n'.join(full_text)
            
        except Exception as e:
            raise Exception(f"Failed to extract text from DOCX bytes: {str(e)}")

    def _extract_excel_from_bytes(self, file_bytes: bytes) -> str:
        """Extract text from Excel bytes"""
        try:
            workbook = openpyxl.load_workbook(io.BytesIO(file_bytes))
            text_parts = []
            
            for sheet_name in workbook.sheetnames:
                worksheet = workbook[sheet_name]
                text_parts.append(f"Sheet: {sheet_name}")
                
                for row in worksheet.iter_rows():
                    row_data = []
                    for cell in row:
                        if cell.value is not None:
                            row_data.append(str(cell.value))
                    if row_data:
                        text_parts.append('\t'.join(row_data))
            
            return '\n'.join(text_parts)
            
        except Exception as e:
            raise Exception(f"Failed to extract text from Excel bytes: {str(e)}")

    def _extract_pdf_from_bytes(self, file_bytes: bytes) -> str:
        """Extract text from PDF bytes"""
        try:
            pdf_reader = PdfReader(io.BytesIO(file_bytes))
            text_parts = []
            
            for page_num, page in enumerate(pdf_reader.pages):
                text = page.extract_text()
                if text.strip():
                    text_parts.append(f"Page {page_num + 1}:")
                    text_parts.append(text)
            
            return '\n'.join(text_parts)
            
        except Exception as e:
            raise Exception(f"Failed to extract text from PDF bytes: {str(e)}")
